from __future__ import annotations

import hashlib
import io
import json
import re
import zipfile
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from .database import Database
from .pii import redact_pii

ALLOWED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}
ALLOWED_ROLES = {"employee", "procurement", "legal", "security", "admin"}
ALLOWED_CLASSIFICATIONS = {"public", "internal", "confidential", "restricted"}
ALLOWED_STATUSES = {"active", "draft", "expired", "superseded"}
MAX_EXTRACTED_CHARS = 2_000_000
MAX_PDF_PAGES = 500
MAX_DOCX_UNCOMPRESSED_BYTES = 30 * 1024 * 1024
MAX_DOCX_ENTRIES = 5000


def parse_front_matter(text: str) -> tuple[dict[str, str], str]:
    if not text.startswith("---\n"):
        return {}, text
    end = text.find("\n---\n", 4)
    if end < 0:
        return {}, text
    block = text[4:end]
    body = text[end + 5 :]
    metadata: dict[str, str] = {}
    for line in block.splitlines():
        if ":" not in line:
            continue
        key, value = line.split(":", 1)
        metadata[key.strip()] = value.strip()
    return metadata, body


def parse_controls(raw: str | None) -> dict[str, str]:
    controls: dict[str, str] = {}
    if not raw:
        return controls
    for item in raw.split(";"):
        if "=" not in item:
            continue
        key, value = item.split("=", 1)
        key = key.strip()
        if key:
            controls[key[:160]] = value.strip()[:500]
    return controls


def extract_text(content: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {suffix or 'none'}")

    if suffix in {".md", ".txt"}:
        text = content.decode("utf-8-sig", errors="replace")
    elif suffix == ".pdf":
        reader = PdfReader(io.BytesIO(content))
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ValueError(f"PDF exceeds the {MAX_PDF_PAGES}-page demonstration limit")
        parts: list[str] = []
        total = 0
        for page in reader.pages:
            part = page.extract_text() or ""
            total += len(part)
            if total > MAX_EXTRACTED_CHARS:
                raise ValueError("Extracted document text exceeds the demonstration limit")
            parts.append(part)
        text = "\n\n".join(parts)
    elif suffix == ".docx":
        try:
            with zipfile.ZipFile(io.BytesIO(content), "r") as archive:
                members = archive.infolist()
                if len(members) > MAX_DOCX_ENTRIES:
                    raise ValueError("DOCX archive contains too many entries")
                if sum(member.file_size for member in members) > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ValueError("DOCX uncompressed size exceeds the demonstration limit")
        except zipfile.BadZipFile as exc:
            raise ValueError("DOCX archive is invalid") from exc
        document = DocxDocument(io.BytesIO(content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(paragraphs)
    else:
        raise ValueError("Unsupported document type")

    if len(text) > MAX_EXTRACTED_CHARS:
        raise ValueError("Extracted document text exceeds the demonstration limit")
    return text


def _clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\t ]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_markdown(text: str, max_chars: int = 1400) -> list[dict[str, str]]:
    current_section = "Document overview"
    buffer: list[str] = []
    chunks: list[dict[str, str]] = []

    def flush() -> None:
        nonlocal buffer
        if not buffer:
            return
        combined = "\n\n".join(buffer).strip()
        while len(combined) > max_chars:
            split_at = combined.rfind(". ", 0, max_chars)
            if split_at < max_chars // 2:
                split_at = max_chars
            part_end = split_at + (1 if split_at < len(combined) and combined[split_at] == "." else 0)
            part = combined[:part_end].strip()
            if not part:
                part = combined[:max_chars]
            chunks.append({"section": current_section, "content": part})
            combined = combined[len(part) :].strip()
        if combined:
            chunks.append({"section": current_section, "content": combined})
        buffer = []

    for raw in text.splitlines():
        line = raw.strip()
        heading = re.match(r"^#{1,4}\s+(.+)$", line)
        if heading:
            flush()
            current_section = heading.group(1).strip()[:300] or "Document section"
            continue
        if not line:
            if buffer and sum(len(item) for item in buffer) >= max_chars * 0.7:
                flush()
            continue
        buffer.append(line)
        if sum(len(item) for item in buffer) >= max_chars:
            flush()
    flush()
    if not chunks and text.strip():
        chunks.append({"section": current_section, "content": text.strip()[:max_chars]})
    return chunks


def _validate_date(value: Any, field_name: str) -> str | None:
    if value in (None, ""):
        return None
    text = str(value).strip()
    try:
        date.fromisoformat(text)
    except ValueError as exc:
        raise ValueError(f"{field_name} must use YYYY-MM-DD format") from exc
    return text


def _validated_roles(raw: Any) -> list[str]:
    roles = list(dict.fromkeys(role.strip().lower() for role in str(raw).split(",") if role.strip()))
    if not roles:
        raise ValueError("At least one allowed role is required")
    unknown = sorted(set(roles) - ALLOWED_ROLES)
    if unknown:
        raise ValueError(f"Unknown allowed role: {', '.join(unknown)}")
    return roles


def _bounded_rank(raw: Any) -> int:
    try:
        rank = int(raw)
    except (TypeError, ValueError) as exc:
        raise ValueError("authority_rank must be an integer") from exc
    if not 0 <= rank <= 100:
        raise ValueError("authority_rank must be between 0 and 100")
    return rank


def make_document_record(
    content: bytes,
    filename: str,
    supplied_metadata: dict[str, Any] | None = None,
    redact_before_index: bool = True,
) -> tuple[dict[str, Any], list[dict[str, str]], dict[str, int]]:
    supplied_metadata = supplied_metadata or {}
    raw_text = extract_text(content, filename)
    front, body = parse_front_matter(raw_text)
    merged = {**front, **{key: value for key, value in supplied_metadata.items() if value not in (None, "")}}
    clean = _clean_text(body)
    if not clean:
        raise ValueError("Document contains no extractable text")

    redaction = redact_pii(clean) if redact_before_index else None
    indexed_text = redaction.text if redaction else clean
    roles = _validated_roles(merged.get("allowed_roles", "employee,procurement,legal,security,admin"))
    classification = str(merged.get("classification", "internal")).strip().lower()
    if classification not in ALLOWED_CLASSIFICATIONS:
        raise ValueError(f"Unsupported classification: {classification}")
    status = str(merged.get("status", "active")).strip().lower()
    if status not in ALLOWED_STATUSES:
        raise ValueError(f"Unsupported document status: {status}")
    title = str(merged.get("title") or Path(filename).stem.replace("_", " ")).strip()
    policy_family = str(merged.get("policy_family", Path(filename).stem.lower().replace(" ", "-"))).strip()
    if not title or not policy_family:
        raise ValueError("Document title and policy family are required")

    document = {
        "title": title[:300],
        "source_filename": filename,
        "content_hash": hashlib.sha256(content).hexdigest(),
        "policy_family": policy_family[:160],
        "department": str(merged.get("department", "General")).strip()[:160] or "General",
        "classification": classification,
        "allowed_roles": roles,
        "effective_date": _validate_date(merged.get("effective_date"), "effective_date"),
        "expires_at": _validate_date(merged.get("expires_at"), "expires_at"),
        "status": status,
        "version": str(merged.get("version", "1.0")).strip()[:80] or "1.0",
        "authority_rank": _bounded_rank(merged.get("authority_rank", 50)),
        "controls": parse_controls(merged.get("controls")),
        "metadata": {
            "description": str(merged.get("description", ""))[:1000],
            "owner": str(merged.get("owner", ""))[:300],
            "redacted_before_index": bool(redact_before_index),
        },
    }
    chunks = chunk_markdown(indexed_text)
    if not chunks:
        raise ValueError("Document produced no indexable chunks")
    counts = redaction.counts if redaction else {}
    return document, chunks, counts



def store_upload_without_overwrite(directory: Path, filename: str, content: bytes) -> Path:
    """Store an upload with exclusive creation so an existing user file is never replaced."""
    directory.mkdir(parents=True, exist_ok=True)
    requested = Path(filename)
    stem = requested.stem[:120] or "uploaded_document"
    suffix = requested.suffix.lower()
    for index in range(1000):
        candidate_name = requested.name if index == 0 else f"{stem}_{index}{suffix}"
        destination = directory / candidate_name
        try:
            with destination.open("xb") as handle:
                handle.write(content)
            return destination
        except FileExistsError:
            continue
        except Exception:
            destination.unlink(missing_ok=True)
            raise
    raise RuntimeError("Upload could not be stored without replacing an existing file")

def _ingest_fingerprint(document: dict[str, Any]) -> str:
    material = {
        key: document.get(key)
        for key in (
            "content_hash", "title", "source_filename", "policy_family", "department",
            "classification", "allowed_roles", "effective_date", "expires_at", "status",
            "version", "authority_rank", "controls",
        )
    }
    material["redacted_before_index"] = bool(document.get("metadata", {}).get("redacted_before_index"))
    encoded = json.dumps(material, sort_keys=True, separators=(",", ":"), ensure_ascii=False).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def ingest_file(
    db: Database,
    path: Path,
    supplied_metadata: dict[str, Any] | None = None,
    redact_before_index: bool = True,
) -> dict[str, Any]:
    content = path.read_bytes()
    document, chunks, counts = make_document_record(
        content,
        path.name,
        supplied_metadata=supplied_metadata,
        redact_before_index=redact_before_index,
    )
    fingerprint = _ingest_fingerprint(document)
    document.setdefault("metadata", {})["ingest_fingerprint"] = fingerprint
    existing = db.find_document_by_source(document["source_filename"], document["policy_family"])
    if (
        existing
        and existing.get("content_hash") == document["content_hash"]
        and existing.get("metadata", {}).get("ingest_fingerprint") == fingerprint
    ):
        return {
            "document_id": existing["id"],
            "title": document["title"],
            "chunks": 0,
            "pii_redactions": counts,
            "skipped_unchanged": True,
            "corpus_generation": db.corpus_generation(),
        }
    document_id = db.upsert_document(document, chunks)
    return {
        "document_id": document_id,
        "title": document["title"],
        "chunks": len(chunks),
        "pii_redactions": counts,
        "skipped_unchanged": False,
        "corpus_generation": db.corpus_generation(),
    }
